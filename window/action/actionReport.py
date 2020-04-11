#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
comment: 计算或软件系统报告

<p>子线程

@author: GanAH  2020/3/9.
@version 1.0.
@contact: dinggan@whu.edu.cn
"""
from PyQt5.QtCore import QThread, pyqtSignal
from docx import Document
from docx.shared import Inches
import os
from win32com import client
from window.file.operationFile import OperationFile
from myConfig.logger import Logger


class Report(QThread):
    overEmit = pyqtSignal()
    exceptionEmit = pyqtSignal(str, str)

    def __init__(self, *args):
        super(Report, self).__init__()
        self.type = args[0]
        self.args = args
        self.logger = Logger().get_logger("REPORT")

    def run(self) -> None:
        if self.type == "C":
            # 坐标转换报告
            self._coorTranReport()

        elif self.type == "L":

            # 徕卡
            if len(self.args) == 6:
                self._leicaReport()

            else:
                print("走歪了")

        else:
            pass

    def _coorTranReport(self):
        filePath = self.args[1]
        coorTranDict = self.args[2]
        paraList = coorTranDict.get("para")
        resultList = self.args[3]

        if filePath != "":
            # 类别区分
            outStrText = []
            if len(paraList) == 4:  # 直接参数法
                outStrText.append("\t========【直接参数转换法】========\n  *版本一测试暂定形式。\n\n  一. 坐标转换参数\n")
                outStrText.append(" 偏移量 Dx/m  偏移量 Dy/m   尺度因子 M/m    旋转角度 SIGMA/°\n")
                outStrText.append(" " + str(paraList))

            elif len(paraList) == 7:  # 最小二乘
                outStrText.append("\t========【最小二乘转换法】========\n  *版本一测试暂定形式。\n\n  一. 坐标转换参数\n")
                outStrText.append(" 偏移量 Dx/m  偏移量 Dy/m   旋转角度THETA/° 尺度因子 M/m    中误差 SIGMA/mm\n")
                outStrText.append(" " + str(paraList))
            else:
                outStrText.append("\t========【正形变换法】========\n  *版本一测试暂定形式。\n\n  一. 坐标转换参数\n")
                outStrText.append(" 公共点及十参数解算结果列表\n")
                outStrText.append(" " + str(paraList))
            # 共性
            outStrText.append("\n\n  二. 坐标转换结果\n")
            outStrText.append("ID    X/m    Y/m  残差TETA_X/m  残差TETA_Y/m\n")
            for i in range(len(resultList)):
                outStrText.append(str(resultList[i][0]) + ", " + str(resultList[i][1]) + ", " + str(
                    resultList[i][2]) + ", " + str(resultList[i][3]) + ", " + str(resultList[i][4]) + "\n")

            # 非直接参数转换详细过程
            if len(coorTranDict) > 2:
                outStrText.append("\n\n  三. 详细解算参数(字典)\n")
                count = 1
                for key in coorTranDict.keys():
                    outStrText.append("  3." + str(count) + " " + key + "\n")
                    count += 1
                    lists = coorTranDict[key]
                    for i in range(len(lists)):
                        outStrText.append("   " + str(lists[i]) + "\n")

            with open(filePath, mode="w", encoding="UTF-8") as f:
                for line in range(len(outStrText)):
                    f.write(outStrText[line])

    def _leicaReport(self):
        filePath = self.args[1]
        stationID = self.args[2]
        stationRemark = self.args[3]
        itemData = self.args[4]
        statusText = self.args[5]
        try:
            # 模板
            docx = Document("./source/template/leica.docx")
            self.tables = docx.tables
            rows = self.tables[0].rows
            allTableLen = len(rows) - 4

            # print("表长", allTableLen, len(rows))
            # 比较数据与表格长度
            if allTableLen < len(itemData):
                self.logger.info("模板表格不足，正在加长表格")
                self._addGroundRow(int((len(itemData) - allTableLen) / 4))
                self.logger.info("表格加长完成,加长范围：" + str(allTableLen) + "~" + str(len(rows)))

            stationIndex_ID = 0
            stationIndex_Remark = 0
            stt = 1

            for i in range(4, len(rows)):
                cols = rows[i].cells
                itemIndex = i - 4
                for k in range(10):
                    if i % 4 == 0:  # 测站与备注信息
                        if k == 0:
                            cell = cols[k]
                            cell.text = stationID[stationIndex_ID]
                            stationIndex_ID += 1
                        elif k == 9:
                            cell = cols[k]
                            if stationIndex_Remark % 2 != 0:
                                cell.text = "测段" + str(stt)
                                stt += 1
                            stationIndex_Remark += 1

                    if k == 1:
                        cell = cols[k]
                        cell.text = "  " + itemData[itemIndex][k - 1]

                    elif k == 3:
                        cell = cols[k]
                        cell.text = "  " + itemData[itemIndex][k - 2]
                    elif k == 5:  # 中间行标记
                        cell = cols[k]
                        cell.text = "  " + itemData[itemIndex][k - 3]
                    elif k > 5 and k < 9:
                        cell = cols[k]
                        cell.text = "  " + itemData[itemIndex][k - 3]

            # 由于表格装不下--- remark.add --
            docx.add_paragraph("测段信息" + "\n" + "* 由于表格太小加入测段说明会拉长不美观，以此代替")
            for i in range(len(stationRemark)):
                docx.add_paragraph(str(stationRemark[i]))
            # 其他详细数据
            paragraph = docx.add_paragraph(statusText)
            # 缩进
            paragraph_format = paragraph.paragraph_format

            paragraph_format.left_indent
            # None
            # 这表示缩进是从样式层次结构中继承的
            paragraph_format.left_indent = Inches(0.3)
            paragraph_format.left_indent
            # 457200
            paragraph_format.left_indent.inches

            dirPath = os.path.dirname(os.path.realpath(filePath))  # 除开文件名的路径
            fileName = (os.path.split(filePath)[1]).split(".")[0]
            # print(fileName)
            # print(dirPath)
            OperationFile().writeTXTFile(statusText, dirPath + "\\" + fileName + ".txt")
            docx.save(filePath)
            self._doc2pdf(filePath, dirPath + "\\" + fileName + ".pdf")

            self.logger.info("导出报告完成:已导出docx-pdf-txt报告组")

        except Exception as e:
            self.logger.info("异常" + e.__str__())

    def _doc2pdf(self, doc_name, pdf_name):
        """
        :word文件转pdf
        :param doc_name word文件名称
        :param pdf_name 转换后pdf文件名称
        """
        try:
            word = client.DispatchEx("Word.Application")
            if os.path.exists(pdf_name):
                os.remove(pdf_name)
            worddoc = word.Documents.Open(doc_name, ReadOnly=1)
            worddoc.SaveAs(pdf_name, FileFormat=17)
            worddoc.Close()
            return pdf_name
        except:
            return 1

    def _addGroundRow(self, groundCount):
        """
        按组添加表格范围
        :param groundCount: 组数，一组四行，容纳 BFFB 模式数据
        :return: None
        """
        for i in range(groundCount):
            # 按组添加
            row_1 = self.tables[0].add_row()
            row_2 = self.tables[0].add_row()
            row_3 = self.tables[0].add_row()
            row_4 = self.tables[0].add_row()
            # 合并首尾项
            row_1.cells[0].text = "走您观🐎"
            # 合并测站区域与备注
            row_1.cells[0].merge(row_4.cells[0])
            row_1.cells[9].merge(row_4.cells[9])
            # 基本区域合并
            for i in range(1, 5, 2):
                row_1.cells[i].merge(row_1.cells[i + 1])
                row_2.cells[i].merge(row_2.cells[i + 1])
                row_3.cells[i].merge(row_3.cells[i + 1])
                row_4.cells[i].merge(row_4.cells[i + 1])

    def killThread(self):
        self.terminate()
